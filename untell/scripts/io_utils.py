"""Input readers — plain text, plus .docx and .pdf when the optional ``[docs]`` extra is installed."""

from __future__ import annotations

import logging

logger = logging.getLogger(__name__)

# Encodings tried in order before falling back to lossy replacement. cp1252 covers the smart quotes,
# em-dashes and ellipses that Word and most Windows editors emit — precisely the characters an
# AI-tells scorer cares about, and precisely what errors="replace" would destroy.
# utf-8-sig FIRST: it decodes plain UTF-8 identically but also strips a leading BOM. With plain
# utf-8 first, a BOM-prefixed file decodes "successfully" and leaves U+FEFF as the first
# character of the text — an invisible char the scorer then sees as content.
_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1252", "latin-1")

# Byte-order marks, longest first — the UTF-32-LE mark (FF FE 00 00) starts with the UTF-16-LE mark
# (FF FE), so checking UTF-16 first would truncate every UTF-32 file to garbage.
#
# Sniffing these is not a nicety. `latin-1` is last in the list above and it maps every one of the
# 256 byte values, so it CANNOT raise UnicodeDecodeError — the loop always "succeeds" there. That
# made the lossy-replacement branch below unreachable dead code, and, far worse, meant a UTF-16
# file (what Windows "Save as -> Unicode" writes) decoded silently into mojibake:
#     'The "smart quotes"'  ->  'ÿþT\x00h\x00e\x00 \x00"\x00s\x00m\x00a\x00r\x00t\x00...'
# and was then scored and rewritten as if that were the user's prose.
_BOMS: tuple[tuple[bytes, str], ...] = (
    (b"\xff\xfe\x00\x00", "utf-32"),
    (b"\x00\x00\xfe\xff", "utf-32"),
    (b"\xef\xbb\xbf", "utf-8-sig"),
    (b"\xff\xfe", "utf-16"),
    (b"\xfe\xff", "utf-16"),
)

# Real text files contain no NUL at all, so a single one means the content is binary or was decoded
# with the wrong codec. Rejecting on any NUL is deliberately strict: the alternative is scoring and
# rewriting mojibake, and this module already prefers a clear error over plausible garbage (see the
# scanned-PDF branch).


def _has_bytes(path: str) -> bool:
    """True when the file has any content at all.

    An empty file and a corrupt one need different sentences: "there is nothing here" is
    actionable, "it may be corrupt or a different format" sends the reader looking for damage that
    is not there.
    """
    import os

    try:
        return os.path.getsize(path) > 0
    except OSError:
        return True  # unreadable size is not evidence of emptiness; let the parser's message stand


def _read_docx(path: str) -> str:
    # Name the extra rather than letting ModuleNotFoundError out. `--file report.docx` is the very
    # first thing a user with a Word document tries, and "No module named 'docx'" does not tell
    # them the fix, or even that there is one — the package is `python-docx` and the import is
    # `docx`, so the error message does not name anything they can install.
    try:
        from docx import Document  # python-docx
    except ModuleNotFoundError:
        raise ValueError(
            f"{path} is a .docx and reading it needs python-docx: pip install 'untell[docs]'"
        ) from None

    # python-docx raises `PackageNotFoundError` for anything that is not a valid OOXML zip — an
    # empty file, a truncated download, a .txt someone renamed. It is not a ValueError or an
    # OSError, so it escaped `read_file_or_exit`'s handlers entirely. MEASURED before this:
    #
    #     untell tells --file zero.docx  ->  exit 1, raw traceback,
    #                                        docx.opc.exceptions.PackageNotFoundError:
    #                                        Package not found at '...zero.docx'
    #
    # and the same for a text file renamed to .docx. "Package not found" is actively misleading:
    # the file is right there and readable, it is simply not a .docx. The convention here is exit 2
    # with one line, which is what `ValueError` gets — the same conversion this function already
    # does for the missing-dependency case a few lines up.
    #
    # Caught around the CONSTRUCTOR only, so a bug in the paragraph/table walk below still raises
    # normally rather than being reported as a bad file.
    try:
        doc = Document(path)
    except Exception as exc:  # PackageNotFoundError and anything else the parser raises
        if not _has_bytes(path):
            raise ValueError(f"{path} is empty, so there is no .docx to read") from None
        raise ValueError(
            f"{path} is not a readable .docx ({type(exc).__name__}). It may be corrupt, "
            f"truncated, or a different format with a .docx extension."
        ) from None

    parts = [p.text for p in doc.paragraphs]
    # Document.paragraphs does NOT descend into tables, so a document whose content lives in a table
    # (very common for reports, forms and CVs) previously read as empty or near-empty and was then
    # scored/humanized as if that text did not exist.
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        raise ValueError(
            f"{path} is a .pdf and reading it needs pypdf: pip install 'untell[docs]'"
        ) from None

    # Same conversion as the .docx path, and for the same reason: pypdf's errors all descend from
    # `PyPdfError`, not from `ValueError` or `OSError`, so every one of them escaped
    # `read_file_or_exit` as a raw traceback with exit 1. MEASURED before this:
    #
    #     zero.pdf       exit 1  pypdf.errors.EmptyFileError: Cannot read an empty file
    #     corrupt.pdf    exit 1  pypdf.errors.PdfStreamError: Stream has ended unexpectedly
    #     encrypted.pdf  exit 1  pypdf.errors.FileNotDecryptedError: File has not been decrypted
    #
    # The encrypted case is the one worth naming separately — the file is perfectly valid and the
    # user needs a password, which is a different action from "your file is broken".
    #
    # Catching by NAME rather than by imported class: pypdf has reorganised this hierarchy across
    # versions, and importing `pypdf.errors` to reference them would add a second import that can
    # itself fail on an older release.
    # Extraction is INSIDE the guard, not just the constructor. `PdfReader(path)` succeeds on an
    # encrypted file and `.pages` gives a list; the refusal only surfaces when a page is actually
    # read. Guarding the constructor alone left that case exiting 1 with a traceback from
    # `page.extract_text()`, which is how this was first "fixed" and then found still broken.
    try:
        pages = PdfReader(path).pages
        texts = [(page.extract_text() or "") for page in pages]
    except Exception as exc:
        name = type(exc).__name__
        if not _has_bytes(path):
            raise ValueError(f"{path} is empty, so there is no .pdf to read") from None
        if "Decrypt" in name or "decrypted" in str(exc):
            raise ValueError(
                f"{path} is password-protected, so its text cannot be read. Decrypt it first "
                f"(e.g. `qpdf --decrypt in.pdf out.pdf`) and retry."
            ) from None
        raise ValueError(
            f"{path} is not a readable .pdf ({name}). It may be corrupt, truncated, or a "
            f"different format with a .pdf extension."
        ) from None

    empty = sum(1 for t in texts if not t.strip())
    if pages and empty == len(pages):
        # Every page yielded nothing: this is a scanned/image PDF and pypdf cannot read it. Returning
        # "" would hand an empty string to the scorer, which would happily report it as clean text.
        raise ValueError(
            f"{path}: no extractable text — this looks like a scanned/image PDF. "
            "Run OCR first (e.g. ocrmypdf) and retry."
        )
    if empty:
        logger.warning(
            "%s: %d of %d pages yielded no text (likely scanned images); the extracted text is "
            "PARTIAL.", path, empty, len(pages),
        )
    return "\n".join(texts)


def _reject_if_binary(path: str, text: str) -> str:
    """Refuse content that decoded "successfully" but is plainly not prose.

    Because latin-1 accepts any byte, a binary file reads back as a string and would be scored as
    text — the detector would report a number for it. Returning garbage is worse than failing.
    """
    if "\x00" in text:
        raise ValueError(
            f"{path}: decoded text contains NUL bytes — this is a binary file, or text in an "
            "encoding this reader could not identify. Convert it to UTF-8 and retry."
        )
    return text


def _read_text(path: str) -> str:
    with open(path, "rb") as fh:
        head = fh.read(4)
    # A byte-order mark is unambiguous: use it rather than guessing. Without this, UTF-16/UTF-32
    # files fell through to latin-1 (which cannot fail) and became mojibake, silently.
    for bom, encoding in _BOMS:
        if head.startswith(bom):
            try:
                with open(path, encoding=encoding) as fh:
                    return _reject_if_binary(path, fh.read())
            except UnicodeDecodeError:
                break  # BOM present but the body does not decode — fall through to the guesses

    for encoding in _TEXT_ENCODINGS:
        try:
            with open(path, encoding=encoding) as fh:
                text = fh.read()
        except UnicodeDecodeError:
            continue
        if encoding == "latin-1":
            # Reached only because every stricter codec failed. latin-1 cannot fail, so arriving
            # here is not evidence the result is right — say so instead of returning it silently.
            logger.warning(
                "%s: decoded as latin-1 only because it maps every byte; %s all failed. If this "
                "file is not Latin-1 the text is mojibake. Convert it to UTF-8 to be sure.",
                path, ", ".join(_TEXT_ENCODINGS[:-1]),
            )
        return _reject_if_binary(path, text)

    # Unreachable while latin-1 is in the list, but kept so removing it degrades safely.
    logger.warning("%s: could not decode with %s; falling back to lossy replacement.",
                   path, ", ".join(_TEXT_ENCODINGS))
    with open(path, encoding="utf-8", errors="replace") as fh:
        return _reject_if_binary(path, fh.read())


def read_file(path: str) -> str:
    """Read text from a file. Handles .docx (python-docx) and .pdf (pypdf) if installed, else text.

    Raises ``ValueError`` with a message meant for a person on the three ways `--file` is usually
    wrong. Before this, each surfaced as a raw traceback from deep in the reader:

        --file no-such-file.txt   FileNotFoundError: [Errno 2] No such file or directory
        --file docs               PermissionError: [Errno 13] Permission denied: 'docs'

    The second one is the reason this exists rather than being left to the OS. Opening a directory
    raises *PermissionError* on Windows, so the message sent the reader to check file permissions
    for a problem that has nothing to do with permissions. `_reject_if_binary` already does this
    job for binary input and says exactly what is wrong and what to do; these three did not.
    """
    import os

    if not os.path.exists(path):
        raise ValueError(f"no such file: {path}")
    if os.path.isdir(path):
        raise ValueError(
            f"{path} is a directory, not a file — pass one document, or use a shell loop"
        )
    if not os.access(path, os.R_OK):
        raise ValueError(f"cannot read {path}: permission denied")

    low = path.lower()
    if low.endswith(".docx"):
        return _read_docx(path)
    if low.endswith(".pdf"):
        return _read_pdf(path)
    return _read_text(path)


def read_file_or_exit(path: str) -> str:
    """``read_file`` for CLIs: a bad path exits 2 with one line, not a traceback.

    Eight commands accept ``--file`` and seven of them let the exception escape, so a typo produced
    a Python stack trace ending in FileNotFoundError, and pointing at a directory produced
    ``PermissionError: [Errno 13] Permission denied`` — a message that sends the reader to check
    permissions for a problem that is not about permissions.

    A stack trace is the right output for a bug in this code and the wrong output for a mistyped
    path. Exit 2 matches argparse's own convention for a usage error, which is what this is.
    """
    import sys

    try:
        return read_file(path)
    except ValueError as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise SystemExit(2) from None
    except OSError as exc:  # anything the explicit guards above did not anticipate
        print(f"error: cannot read {path}: {exc.strerror or exc}", file=sys.stderr)
        raise SystemExit(2) from None


def read_stdin_or_none() -> str | None:
    """Read piped stdin, or return None when stdin is an interactive terminal.

    `sys.stdin.read()` on a TTY blocks until the user sends EOF, and none of these commands prints
    a prompt first — so `untell tells` typed with no argument produced no output and no error and
    looked hung, when the answer the user wanted was the usage line. MEASURED by modelling a
    terminal (isatty True, read() raising instead of blocking): 6 of the 7 stdin-reading commands
    blocked; `scrub` alone guarded and returned 2.

    This is that guard, hoisted so the seventh copy is a call rather than a copy. Returning None
    rather than raising keeps the caller's own "no input" message and exit code, which the tests in
    this repo assert per command.
    """
    import sys

    try:
        interactive = sys.stdin.isatty()
    except Exception:
        # A replaced or closed stream in a test harness has no isatty; treat it as non-interactive
        # so piped and captured input still reaches the command.
        interactive = False
    if interactive:
        return None
    try:
        text = sys.stdin.read()
    except UnicodeDecodeError:
        # Binary/undecodable stdin (a piped binary file, curl output). The command's contract is
        # a clean "no input" exit 2, not a Python traceback. Same path as empty stdin. MEASURED:
        # b'\x00\x01\x02\xff' piped to untell-score leaked UnicodeDecodeError before this guard.
        return None
    if "\x00" in text:
        # Same policy as `_reject_if_binary` for files: real text contains no NUL at all, so one
        # means the content is binary. This guard exists because some codecs (Windows cp1252 /
        # latin-1 pipes) decode NUL bytes instead of raising, which let a piped binary file
        # through to be scored and "humanized" as prose with exit 0. MEASURED on this host:
        # `untell humanize` on piped b'\x00\x01\x02\xff' reported "humanization complete".
        # Callers already turn None into their clean "no input" exit 2.
        return None
    return text


def configure_utf8_io() -> None:
    """Force UTF-8 on stdin/stdout/stderr so non-ASCII text never crashes a Windows console.

    Windows defaults the standard streams to the locale code page (commonly cp1252); piping or
    printing emoji / accented / RTL text then raises ``UnicodeDecodeError``/``UnicodeEncodeError``.
    Each stream is reconfigured independently and best-effort (a replaced stream in tests may lack
    ``reconfigure`` — that is fine, it is skipped).

    ``errors="replace"`` as well as the encoding, so a LONE SURROGATE in the text (invalid Unicode
    that reaches a CLI via argv or a broken file decode) prints as U+FFFD instead of killing the
    process with UnicodeEncodeError. Fuzz-found: `untell scrub` printing a surrogate argv crashed
    with a traceback; argparse's own stderr path did the same on a surrogate-bearing --tier value.
    """
    import sys

    for stream in (sys.stdin, sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass
